import os
import io
import re
import hashlib
import mimetypes
from datetime import datetime, date
from pathlib import PurePosixPath

import fitz
import pandas as pd
import streamlit as st
from docx import Document
from openai import OpenAI
from supabase import create_client


# ============================================================
# MELO ALVES — GOVERNANÇA TRIBUTÁRIA
# Streamlit + Supabase + OpenAI
# ============================================================

APP_NAME = "Melo Alves — Governança Tributária"

st.set_page_config(
    page_title=APP_NAME,
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
    <style>
    .block-container {
        padding-top: 1.25rem;
        padding-bottom: 3rem;
    }

    .ma-card {
        border: 1px solid #e5e7eb;
        border-radius: 14px;
        padding: 16px;
        background: white;
        margin-bottom: 12px;
    }

    .ma-muted {
        color: #667085;
        font-size: .90rem;
    }

    div[data-testid="stMetric"] {
        border: 1px solid #eaecf0;
        border-radius: 14px;
        padding: 12px 14px;
        background: white;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# CONFIGURAÇÕES / SECRETS
# ============================================================

def secret(name, default=""):
    try:
        if name in st.secrets:
            return str(st.secrets[name])
    except Exception:
        pass

    return os.getenv(name, default)


OPENAI_API_KEY = secret("OPENAI_API_KEY")

# Pode depois adicionar OPENAI_MODEL nos Secrets.
# GPT-5.6 será tentado primeiro.
OPENAI_MODEL = secret("OPENAI_MODEL", "gpt-5.6")

SUPABASE_URL = secret("SUPABASE_URL")
SUPABASE_SECRET_KEY = secret("SUPABASE_SECRET_KEY")

CLIENT_BUCKET = secret(
    "SUPABASE_CLIENT_BUCKET",
    "client-documents"
)

LEGAL_BUCKET = secret(
    "SUPABASE_LEGAL_BUCKET",
    "legal-library"
)

# Login temporário do MVP.
# Depois configuraremos autenticação profissional.
ADMIN_EMAIL = secret(
    "ADMIN_EMAIL",
    "admin@meloalves.local"
)

ADMIN_PASSWORD = secret(
    "ADMIN_PASSWORD",
    "Admin@123"
)


# ============================================================
# CONEXÕES
# ============================================================

@st.cache_resource
def sb_client():

    if not SUPABASE_URL:
        raise RuntimeError(
            "SUPABASE_URL não configurada."
        )

    if not SUPABASE_SECRET_KEY:
        raise RuntimeError(
            "SUPABASE_SECRET_KEY não configurada."
        )

    return create_client(
        SUPABASE_URL,
        SUPABASE_SECRET_KEY
    )


@st.cache_resource
def oa_client():

    if not OPENAI_API_KEY:
        raise RuntimeError(
            "OPENAI_API_KEY não configurada."
        )

    return OpenAI(
        api_key=OPENAI_API_KEY
    )


# ============================================================
# FUNÇÕES GERAIS
# ============================================================

def now_iso():

    return datetime.now().astimezone().isoformat()


def clean_name(name):

    name = re.sub(
        r"[^A-Za-z0-9._-]+",
        "_",
        name or "arquivo"
    )

    return name[:180]


def row_name(row):

    return (
        row.get("trade_name")
        or row.get("legal_name")
        or row.get("name")
        or f"Cliente {row.get('id', '')}"
    )


def get_value(row, *keys, default=""):

    if not row:
        return default

    for key in keys:

        value = row.get(key)

        if value not in (None, ""):
            return value

    return default


def format_money(value):

    try:

        number = float(value or 0)

        return (
            f"R$ {number:,.2f}"
            .replace(",", "X")
            .replace(".", ",")
            .replace("X", ".")
        )

    except Exception:

        return "R$ 0,00"


# ============================================================
# BANCO SUPABASE
# ============================================================

MISSING_COLUMN = re.compile(
    r"Could not find the '([^']+)' column",
    re.I
)


def db_select(
    table,
    filters=None,
    order=None,
    desc=True,
    limit=None,
    silent=False
):

    try:

        query = (
            sb_client()
            .table(table)
            .select("*")
        )

        for column, value in (filters or {}).items():

            query = query.eq(
                column,
                value
            )

        if order:

            query = query.order(
                order,
                desc=desc
            )

        if limit:

            query = query.limit(limit)

        result = query.execute()

        return result.data or []

    except Exception as error:

        if not silent:

            st.error(
                f"Erro ao consultar `{table}`: {error}"
            )

        return []


def db_table_exists(table):

    try:

        (
            sb_client()
            .table(table)
            .select("*")
            .limit(1)
            .execute()
        )

        return True

    except Exception:

        return False


def first_existing_table(*names):

    for name in names:

        if db_table_exists(name):
            return name

    return None


def remove_missing_column(payload, error):

    match = MISSING_COLUMN.search(
        str(error)
    )

    if match:

        column = match.group(1)

        if column in payload:

            new_payload = dict(payload)

            new_payload.pop(
                column,
                None
            )

            return new_payload, True

    return payload, False


def db_insert_flexible(
    table,
    payload,
    silent=False
):

    data = {
        key: value
        for key, value in payload.items()
        if value is not None
    }

    for _ in range(30):

        try:

            result = (
                sb_client()
                .table(table)
                .insert(data)
                .execute()
            )

            rows = result.data or []

            if rows:
                return rows[0]

            return data

        except Exception as error:

            new_data, changed = remove_missing_column(
                data,
                error
            )

            if changed:

                data = new_data
                continue

            if not silent:

                st.error(
                    f"Não foi possível salvar em "
                    f"`{table}`: {error}"
                )

            return None

    return None


def db_update_flexible(
    table,
    row_id,
    payload,
    silent=False
):

    data = {
        key: value
        for key, value in payload.items()
        if value is not None
    }

    for _ in range(30):

        try:

            result = (
                sb_client()
                .table(table)
                .update(data)
                .eq("id", row_id)
                .execute()
            )

            rows = result.data or []

            if rows:
                return rows[0]

            return data

        except Exception as error:

            new_data, changed = remove_missing_column(
                data,
                error
            )

            if changed:

                data = new_data
                continue

            if not silent:

                st.error(
                    f"Não foi possível atualizar "
                    f"`{table}`: {error}"
                )

            return None

    return None


def select_by_id(
    table,
    row_id
):

    rows = db_select(
        table,
        {"id": row_id},
        silent=True
    )

    return rows[0] if rows else None


# ============================================================
# SUPABASE STORAGE
# ============================================================

def storage_upload(
    bucket,
    path,
    data,
    content_type=None
):

    content_type = (
        content_type
        or mimetypes.guess_type(path)[0]
        or "application/octet-stream"
    )

    try:

        (
            sb_client()
            .storage
            .from_(bucket)
            .upload(
                path=path,
                file=data,
                file_options={
                    "content-type": content_type,
                    "upsert": "false"
                }
            )
        )

        return True, None

    except Exception as error:

        return False, str(error)


def storage_download(
    bucket,
    path
):

    try:

        data = (
            sb_client()
            .storage
            .from_(bucket)
            .download(path)
        )

        return data, None

    except Exception as error:

        return None, str(error)


def storage_list(
    bucket,
    folder=""
):

    try:

        return (
            sb_client()
            .storage
            .from_(bucket)
            .list(folder)
            or []
        )

    except Exception:

        return []


# ============================================================
# LEITURA DOS ARQUIVOS
# ============================================================

def extract_text(
    filename,
    raw
):

    filename = filename or ""

    extension = ""

    if "." in filename:

        extension = (
            filename
            .lower()
            .rsplit(".", 1)[-1]
        )

    try:

        # PDF
        if extension == "pdf":

            document = fitz.open(
                stream=raw,
                filetype="pdf"
            )

            parts = []

            for page_number, page in enumerate(document):

                text = (
                    page.get_text("text")
                    or ""
                )

                if text.strip():

                    parts.append(
                        f"\n--- PÁGINA "
                        f"{page_number + 1} ---\n"
                        f"{text}"
                    )

            return "".join(parts)[:600000]

        # DOCX
        if extension == "docx":

            document = Document(
                io.BytesIO(raw)
            )

            return "\n".join(
                paragraph.text
                for paragraph in document.paragraphs
            )[:600000]

        # Texto
        if extension in (
            "txt",
            "md",
            "csv",
            "json",
            "xml"
        ):

            return raw.decode(
                "utf-8",
                errors="ignore"
            )[:600000]

        # Excel
        if extension in (
            "xlsx",
            "xls"
        ):

            bio = io.BytesIO(raw)

            sheets = pd.read_excel(
                bio,
                sheet_name=None
            )

            output = []

            for sheet_name, df in sheets.items():

                output.append(
                    f"\n--- PLANILHA: "
                    f"{sheet_name} ---\n"
                )

                output.append(
                    df
                    .fillna("")
                    .astype(str)
                    .to_csv(index=False)
                )

            return "".join(output)[:600000]

    except Exception as error:

        return (
            "[Falha na extração automática: "
            f"{error}]"
        )

    return ""


def sidecar_path(
    original_path
):

    path = PurePosixPath(
        original_path
    )

    stem = (
        path.name
        .rsplit(".", 1)[0]
    )

    return str(
        path.parent
        / "_text"
        / f"{stem}.txt"
    )


# ============================================================
# STATUS
# ============================================================

COMMERCIAL_STATUSES = [

    "Em prospecção",
    "Diagnóstico em andamento",
    "Proposta enviada",
    "Não contratou",
    "Contratado"

]


OPERATIONAL_STATUSES = [

    "Aguardando contratação",
    "Aguardando início da execução",
    "Em execução",
    "Executado",
    "Finalizado",
    "Pausado"

]


CASE_STATUSES = [

    "Abertura",
    "Aguardando documentos",
    "Análise preliminar",
    "Diagnóstico",
    "Proposta",
    "Contratado",
    "Em execução",
    "Concluído",
    "Arquivado"

]


# ============================================================
# CHECKLIST DOCUMENTAL
# ============================================================

CHECKLIST = [

    (
        "relatorio_fiscal",
        "Relatório de Situação Fiscal atualizado",
        True,
        "Receita Federal/PGFN completo e recente."
    ),

    (
        "cnd",
        "CND/CPEN e demais certidões",
        True,
        "Federal, estadual, municipal, "
        "FGTS e trabalhista, conforme o caso."
    ),

    (
        "societario",
        "Contrato social, alterações e procurações",
        True,
        "Documentos societários e de representação."
    ),

    (
        "ecf",
        "ECF dos últimos 5 anos",
        True,
        "Arquivo, recibo e apurações de "
        "IRPJ/CSLL, quando aplicável."
    ),

    (
        "ecd",
        "ECD, balancetes e razão dos últimos 5 anos",
        True,
        "Conforme regime e pertinência."
    ),

    (
        "dctf",
        "DCTF/DCTFWeb e recibos",
        True,
        "Declarações e respectivos recibos."
    ),

    (
        "pagamentos",
        "DARFs, pagamentos, parcelamentos e extratos",
        True,
        "Comprovantes e extratos de conta fiscal."
    ),

    (
        "fontes",
        "Fontes pagadoras e retenções",
        False,
        "DIRF/EFD-Reinf e relatórios do e-CAC, "
        "quando houver."
    ),

    (
        "ato",
        "Auto de infração, intimação, despacho ou decisão",
        False,
        "Documento integral e prova da ciência/data."
    ),

    (
        "proc_adm",
        "Processo administrativo completo",
        False,
        "Capa a capa, com decisões, recursos, "
        "manifestações e anexos."
    ),

    (
        "proc_jud",
        "Processo judicial completo",
        False,
        "Capa a capa, incluindo decisões "
        "e certidões relevantes."
    ),

    (
        "perdcomp",
        "PER/DCOMP, despachos e comunicações",
        False,
        "Pedidos, declarações de compensação, "
        "despachos decisórios e ciência."
    ),

    (
        "contratos",
        "Editais e contratos públicos/privados",
        False,
        "Edital, proposta, planilha, contrato, "
        "aditivos, matriz de riscos e execução."
    ),

    (
        "creditos",
        "Memória de cálculo e planilhas de créditos",
        False,
        "Origem, período, base, valor, utilização e saldo."
    ),

    (
        "outros",
        "Outros documentos relevantes",
        False,
        "Qualquer prova adicional relacionada "
        "à dor do cliente."
    )

]


# ============================================================
# CLIENTES / CASOS
# ============================================================

def clients():

    return db_select(
        "clients",
        order="created_at",
        desc=True,
        silent=True
    )


def cases():

    return db_select(
        "cases",
        order="created_at",
        desc=True,
        silent=True
    )


def client_map():

    return {
        str(client.get("id")): client
        for client in clients()
    }


def get_client(
    client_id
):

    return select_by_id(
        "clients",
        client_id
    )


def get_case(
    case_id
):

    return select_by_id(
        "cases",
        case_id
    )


def case_documents(
    case_id
):

    rows = db_select(
        "documents",
        silent=True
    )

    return [
        row
        for row in rows
        if str(row.get("case_id", ""))
        == str(case_id)
    ]


def document_path(row):

    return get_value(
        row,
        "storage_path",
        "stored_path",
        "path"
    )


def document_name(row):

    return get_value(
        row,
        "original_name",
        "file_name",
        "name",
        default="Documento"
    )


# ============================================================
# HISTÓRICO DO CLIENTE
# ============================================================

def add_timeline(
    client_id,
    title,
    details="",
    event_type="Atualização"
):

    table = first_existing_table(
        "client_timeline",
        "timeline"
    )

    if not table:
        return

    db_insert_flexible(
        table,
        {
            "client_id": client_id,

            "event_type": event_type,
            "type": event_type,

            "title": title,

            "description": details,
            "details": details,

            "created_at": now_iso()
        },
        silent=True
    )


# ============================================================
# CONTEXTO DOS DOCUMENTOS
# ============================================================

def read_sidecar_for_document(row):

    path = document_path(row)

    extracted = get_value(
        row,
        "extracted_text",
        default=""
    )

    if extracted:
        return extracted

    if not path:
        return ""

    raw, error = storage_download(
        CLIENT_BUCKET,
        sidecar_path(path)
    )

    if raw:

        return raw.decode(
            "utf-8",
            errors="ignore"
        )

    return ""


def build_case_context(
    case_id,
    max_chars=120000
):

    documents = case_documents(
        case_id
    )

    blocks = []

    used = 0

    for document in documents:

        text = read_sidecar_for_document(
            document
        )

        if not text:
            continue

        chunk = text[:35000]

        header = (
            "\n\n===== DOCUMENTO: "
            f"{document_name(document)}"
            " | CATEGORIA: "
            f"{get_value(document, 'category', 'label', 'rkey')}"
            " =====\n"
        )

        block = header + chunk

        if used + len(block) > max_chars:

            block = block[
                : max_chars - used
            ]

        blocks.append(block)

        used += len(block)

        if used >= max_chars:
            break

    return "".join(blocks)


# ============================================================
# BASE JURÍDICA
# ============================================================

def load_legal_context(
    max_chars=70000
):

    files = storage_list(
        LEGAL_BUCKET,
        "_text"
    )

    blocks = []

    used = 0

    for item in reversed(files[-30:]):

        name = item.get("name")

        if not name:
            continue

        if not name.endswith(".txt"):
            continue

        raw, error = storage_download(
            LEGAL_BUCKET,
            f"_text/{name}"
        )

        if not raw:
            continue

        text = raw.decode(
            "utf-8",
            errors="ignore"
        )

        block = (
            "\n\n===== BASE JURÍDICA: "
            f"{name} =====\n"
            f"{text[:25000]}"
        )

        if used + len(block) > max_chars:

            block = block[
                : max_chars - used
            ]

        blocks.append(block)

        used += len(block)

        if used >= max_chars:
            break

    return "".join(blocks)


# ============================================================
# OPENAI
# ============================================================

SYSTEM_LEGAL = """

Você é o motor técnico de apoio do
Melo Alves — Governança Tributária.

Atue como assistente jurídico-tributário brasileiro
para apoio de advogado humano.

REGRAS OBRIGATÓRIAS:

1. Nunca invente documento, fato, prazo, dispositivo
legal, precedente, valor ou conclusão.

2. Separe claramente:
- fatos comprovados;
- informações relatadas;
- inferências;
- pontos pendentes.

3. Cite nominalmente os documentos utilizados quando
eles forem fornecidos no contexto.

4. Quando utilizar pesquisa web, priorize fontes oficiais
brasileiras:
- Planalto;
- Receita Federal;
- PGFN;
- CGIBS;
- Diário Oficial;
- STF;
- STJ;
- CARF;
- tribunais;
- Secretarias de Fazenda.

5. Informe quando norma ou entendimento precisar
de confirmação de vigência.

6. Não prometa:
- êxito;
- liminar;
- crédito;
- economia;
- recuperação;
- resultado administrativo ou judicial.

7. Toda peça, diagnóstico e estratégia são minutas
para revisão e aprovação do advogado responsável.

8. Em matéria tributária, trate como pontos críticos:
- prazo;
- decadência;
- prescrição;
- competência;
- autoridade;
- legitimidade;
- prova documental.

9. Para recuperação de créditos, diferencie:
- crédito potencial;
- crédito identificado;
- crédito documentalmente validado;
- crédito juridicamente aproveitável;
- benefício econômico efetivamente realizado.

10. Responda sempre em português do Brasil,
de maneira técnica, organizada e executável.

"""


def openai_response(
    prompt,
    web_search=False
):

    client = oa_client()

    tools = None

    if web_search:

        tools = [
            {
                "type": "web_search"
            }
        ]

    models = [
        OPENAI_MODEL
    ]

    # Fallback caso o modelo principal
    # não esteja liberado na conta.
    if OPENAI_MODEL != "gpt-5-mini":

        models.append(
            "gpt-5-mini"
        )

    last_error = None

    for model in models:

        try:

            arguments = {

                "model": model,

                "instructions": SYSTEM_LEGAL,

                "input": prompt

            }

            if tools:

                arguments["tools"] = tools

            response = (
                client
                .responses
                .create(**arguments)
            )

            return (
                response.output_text,
                model
            )

        except Exception as error:

            last_error = error

    raise RuntimeError(
        str(last_error)
    )


# ============================================================
# SALVAR ANÁLISES
# ============================================================

def save_case_ai(
    case_id,
    client_id,
    kind,
    text,
    model
):

    if kind == "preliminary":

        field = "ai_preliminary"
        status = "Análise preliminar"
        title = "Análise preliminar"

    else:

        field = "ai_diagnosis"
        status = "Diagnóstico"
        title = "Diagnóstico tributário"

    db_update_flexible(
        "cases",
        case_id,
        {
            field: text,
            "status": status,
            "updated_at": now_iso()
        },
        silent=True
    )

    table = first_existing_table(

        "diagnostics",
        "case_diagnostics",
        "generated_documents",
        "generated"

    )

    if table:

        db_insert_flexible(
            table,
            {
                "client_id": client_id,

                "case_id": case_id,

                "diagnostic_type": kind,

                "doc_type": kind,

                "title": title,

                "content": text,

                "model": model,

                "created_at": now_iso()
            },
            silent=True
        )

    add_timeline(
        client_id,
        f"{title} gerado",
        f"Modelo utilizado: {model}",
        "Inteligência Artificial"
    )


# ============================================================
# GERAR DOCX
# ============================================================

def text_to_docx(
    title,
    text
):

    output = io.BytesIO()

    document = Document()

    document.add_heading(
        title,
        0
    )

    document.add_paragraph(
        "Melo Alves — Governança Tributária"
    )

    document.add_paragraph(
        "Gerado em: "
        + datetime.now().strftime(
            "%d/%m/%Y %H:%M"
        )
    )

    document.add_paragraph(
        "Minuta gerada com apoio de inteligência "
        "artificial. Revisão jurídica humana obrigatória "
        "antes de assinatura, protocolo ou envio."
    )

    for line in (text or "").splitlines():

        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("### "):

            document.add_heading(
                stripped[4:],
                level=3
            )

        elif stripped.startswith("## "):

            document.add_heading(
                stripped[3:],
                level=2
            )

        elif stripped.startswith("# "):

            document.add_heading(
                stripped[2:],
                level=1
            )

        elif stripped.startswith("- "):

            document.add_paragraph(
                stripped[2:],
                style="List Bullet"
            )

        elif stripped.startswith("* "):

            document.add_paragraph(
                stripped[2:],
                style="List Bullet"
            )

        else:

            document.add_paragraph(
                stripped
            )

    document.save(output)

    output.seek(0)

    return output.getvalue()


# ============================================================
# LOGIN
# ============================================================

def login_page():

    left, center, right = st.columns(
        [1, 1.2, 1]
    )

    with center:

        st.markdown(
            "## ⚖️ Melo Alves"
        )

        st.caption(
            "Governança Tributária — acesso restrito"
        )

        with st.form(
            "login_form"
        ):

            email = st.text_input(
                "E-mail",
                value=ADMIN_EMAIL
            )

            password = st.text_input(
                "Senha",
                type="password"
            )

            submitted = (
                st.form_submit_button(
                    "Entrar",
                    width="stretch"
                )
            )

        if submitted:

            if (
                email.strip().lower()
                == ADMIN_EMAIL.lower()
                and password
                == ADMIN_PASSWORD
            ):

                st.session_state[
                    "user"
                ] = {

                    "name":
                    "Administrador Melo Alves",

                    "email":
                    email.strip(),

                    "role":
                    "Administrador"

                }

                st.rerun()

            else:

                st.error(
                    "E-mail ou senha inválidos."
                )


if "user" not in st.session_state:

    login_page()

    st.stop()


# ============================================================
# SIDEBAR
# ============================================================

with st.sidebar:

    st.markdown(
        "## ⚖️ MELO ALVES"
    )

    st.caption(
        "Governança Tributária"
    )

    st.markdown(
        f"**{st.session_state['user']['name']}**"
    )

    st.divider()

    page = st.radio(
        "Navegação",
        [

            "Dashboard",

            "Novo atendimento",

            "Clientes",

            "Casos",

            "Tarefas e prazos",

            "Documentos",

            "Base jurídica",

            "Propostas e contratos",

            "Relatórios",

            "Configurações"

        ],
        label_visibility="collapsed"
    )

    st.divider()

    if st.button(
        "Sair",
        width="stretch"
    ):

        st.session_state.clear()

        st.rerun()


# ============================================================
# INTERFACE AUXILIAR
# ============================================================

def header(
    title,
    subtitle=""
):

    st.markdown(
        f"# {title}"
    )

    if subtitle:

        st.caption(
            subtitle
        )

    st.divider()


def choose_client(
    label="Cliente",
    key=None
):

    rows = clients()

    if not rows:

        st.warning(
            "Cadastre um cliente primeiro."
        )

        return None

    options = {
        str(row["id"]): row
        for row in rows
    }

    selected = st.selectbox(

        label,

        list(options.keys()),

        format_func=lambda value:
        (
            f"{row_name(options[value])}"
            " — "
            f"{get_value(options[value], 'cnpj', default='sem CNPJ')}"
        ),

        key=key
    )

    return options[selected]


def choose_case(
    label="Caso",
    key=None,
    client_id=None
):

    rows = cases()

    if client_id is not None:

        rows = [

            row
            for row in rows

            if str(
                row.get("client_id")
            )
            == str(client_id)

        ]

    if not rows:

        st.warning(
            "Nenhum caso cadastrado."
        )

        return None

    options = {
        str(row["id"]): row
        for row in rows
    }

    selected = st.selectbox(

        label,

        list(options.keys()),

        format_func=lambda value:
        (
            f"#{str(value)[:8]}"
            " — "
            f"{get_value(options[value], 'title', default='Caso')}"
        ),

        key=key
    )

    return options[selected]


# ============================================================
# DASHBOARD
# ============================================================

if page == "Dashboard":

    header(
        "Dashboard",
        "Visão consolidada de clientes, casos, "
        "execução e documentos."
    )

    client_rows = clients()

    case_rows = cases()

    document_rows = db_select(
        "documents",
        silent=True
    )

    task_rows = db_select(
        "tasks",
        silent=True
    )

    col1, col2, col3, col4 = st.columns(4)

    col1.metric(
        "Clientes",
        len(client_rows)
    )

    col2.metric(
        "Casos",
        len(case_rows)
    )

    col3.metric(
        "Documentos",
        len(document_rows)
    )

    col4.metric(
        "Tarefas",
        len(task_rows)
    )


    # FUNIL COMERCIAL

    status_counts = {}

    for client in client_rows:

        status = get_value(

            client,

            "commercial_status",

            "status",

            default="Sem status"

        )

        status_counts[status] = (
            status_counts.get(
                status,
                0
            )
            + 1
        )


    if status_counts:

        st.subheader(
            "Funil comercial"
        )

        dataframe = pd.DataFrame([

            {
                "Status": status,
                "Clientes": amount
            }

            for status, amount
            in status_counts.items()

        ])

        st.dataframe(
            dataframe,
            hide_index=True,
            width="stretch"
        )


    st.subheader(
        "Clientes recentes"
    )

    if client_rows:

        rows = []

        for client in client_rows[:12]:

            rows.append({

                "Cliente":
                row_name(client),

                "CNPJ":
                get_value(
                    client,
                    "cnpj"
                ),

                "Regime":
                get_value(
                    client,
                    "tax_regime"
                ),

                "Status comercial":
                get_value(
                    client,
                    "commercial_status",
                    "status"
                ),

                "Status operacional":
                get_value(
                    client,
                    "operational_status"
                )

            })

        st.dataframe(
            pd.DataFrame(rows),
            hide_index=True,
            width="stretch"
        )

    else:

        st.info(
            "Nenhum cliente cadastrado ainda."
        )


# ============================================================
# NOVO ATENDIMENTO
# ============================================================

elif page == "Novo atendimento":

    header(
        "Novo atendimento",
        "Cadastre a empresa, abra o caso, "
        "anexe os documentos e descreva a dor."
    )


    # --------------------------------------------------------
    # CLIENTE
    # --------------------------------------------------------

    st.subheader(
        "1. Cadastrar cliente"
    )

    with st.form(
        "new_client"
    ):

        column1, column2 = (
            st.columns(2)
        )

        legal_name = (
            column1.text_input(
                "Razão social *"
            )
        )

        trade_name = (
            column2.text_input(
                "Nome fantasia"
            )
        )

        cnpj = (
            column1.text_input(
                "CNPJ"
            )
        )

        tax_regime = (
            column2.selectbox(
                "Regime tributário",
                [
                    "Simples Nacional",
                    "Lucro Presumido",
                    "Lucro Real",
                    "Outro"
                ]
            )
        )

        annual_revenue = (
            column1.number_input(
                "Faturamento anual aproximado",
                min_value=0.0,
                step=10000.0
            )
        )

        sector = (
            column2.text_input(
                "Setor/atividade"
            )
        )

        contact_name = (
            column1.text_input(
                "Contato principal"
            )
        )

        contact_email = (
            column2.text_input(
                "E-mail"
            )
        )

        contact_phone = (
            column1.text_input(
                "Telefone/WhatsApp"
            )
        )

        accountant = (
            column2.text_input(
                "Contador/escritório contábil"
            )
        )

        submit_client = (
            st.form_submit_button(
                "Criar cliente",
                width="stretch"
            )
        )


    if submit_client:

        if not legal_name.strip():

            st.error(
                "Informe a razão social."
            )

        else:

            existing = [

                client
                for client in clients()

                if cnpj
                and get_value(
                    client,
                    "cnpj"
                ) == cnpj

            ]

            if existing:

                st.warning(
                    "Já existe cliente com este CNPJ."
                )

            else:

                created = db_insert_flexible(

                    "clients",

                    {

                        "legal_name":
                        legal_name.strip(),

                        "trade_name":
                        trade_name.strip(),

                        "cnpj":
                        cnpj.strip(),

                        "tax_regime":
                        tax_regime,

                        "annual_revenue":
                        annual_revenue,

                        "sector":
                        sector.strip(),

                        "accountant":
                        accountant.strip(),

                        "contact_name":
                        contact_name.strip(),

                        "contact_email":
                        contact_email.strip(),

                        "contact_phone":
                        contact_phone.strip(),

                        "commercial_status":
                        "Em prospecção",

                        "operational_status":
                        "Aguardando contratação",

                        "status":
                        "Ativo",

                        "created_at":
                        now_iso(),

                        "updated_at":
                        now_iso()

                    }

                )

                if created:

                    st.success(
                        "Cliente criado no Supabase."
                    )

                    add_timeline(

                        created.get("id"),

                        "Cliente cadastrado",

                        "Novo atendimento iniciado.",

                        "Cadastro"

                    )

                    st.rerun()


    # --------------------------------------------------------
    # CASO
    # --------------------------------------------------------

    st.divider()

    st.subheader(
        "2. Abrir caso"
    )

    client = choose_client(
        "Selecione o cliente",
        key="new_case_client"
    )


    if client:

        with st.form(
            "new_case"
        ):

            column1, column2 = (
                st.columns(2)
            )

            title = (
                column1.text_input(
                    "Título do caso *",
                    placeholder=(
                        "Ex.: PER/DCOMP "
                        "não homologado"
                    )
                )
            )

            case_type = (
                column2.selectbox(
                    "Tipo",
                    [

                        "Diagnóstico tributário",

                        "PER/DCOMP / compensação",

                        "Auto de infração",

                        "Regularidade fiscal / certidão",

                        "Recuperação de créditos",

                        "Reequilíbrio de contrato público",

                        "Consulta tributária",

                        "Contencioso judicial",

                        "Outro"

                    ]
                )
            )

            urgency = (
                column1.selectbox(
                    "Urgência",
                    [
                        "Normal",
                        "Alta",
                        "Crítica"
                    ]
                )
            )

            pain_summary = (
                column2.text_input(
                    "Resumo da dor"
                )
            )

            pain_details = (
                st.text_area(
                    "Relato completo do cliente",
                    height=180
                )
            )

            objectives = (
                st.text_area(
                    "Objetivo esperado",
                    height=90
                )
            )

            submit_case = (
                st.form_submit_button(
                    "Criar caso",
                    width="stretch"
                )
            )


        if submit_case:

            if not title.strip():

                st.error(
                    "Informe o título."
                )

            else:

                created = db_insert_flexible(

                    "cases",

                    {

                        "client_id":
                        client.get("id"),

                        "title":
                        title.strip(),

                        "case_type":
                        case_type,

                        "urgency":
                        urgency,

                        "status":
                        "Aguardando documentos",

                        "pain_summary":
                        pain_summary.strip(),

                        "pain_details":
                        pain_details.strip(),

                        "objectives":
                        objectives.strip(),

                        "created_at":
                        now_iso(),

                        "updated_at":
                        now_iso()

                    }

                )

                if created:

                    add_timeline(

                        client.get("id"),

                        "Caso aberto",

                        title,

                        "Caso"

                    )

                    st.success(
                        "Caso criado. Vá para "
                        "a página Casos para "
                        "anexar documentos."
                    )


# ============================================================
# CLIENTES
# ============================================================

elif page == "Clientes":

    header(
        "Clientes",
        "CRM tributário: dados, status, "
        "casos e histórico da empresa."
    )

    client_rows = clients()


    if not client_rows:

        st.info(
            "Nenhum cliente cadastrado."
        )

    else:

        selected = choose_client(
            "Abrir ficha do cliente",
            key="client_ficha"
        )


        if selected:

            client_id = (
                selected.get("id")
            )

            st.markdown(
                f"## {row_name(selected)}"
            )

            st.caption(
                "CNPJ: "
                f"{get_value(selected, 'cnpj', default='-')}"
                " | Regime: "
                f"{get_value(selected, 'tax_regime', default='-')}"
            )


            tab1, tab2, tab3, tab4 = (
                st.tabs(
                    [
                        "Visão geral",
                        "Status",
                        "Casos",
                        "Histórico"
                    ]
                )
            )


            # VISÃO GERAL

            with tab1:

                col1, col2, col3 = (
                    st.columns(3)
                )

                col1.metric(
                    "Faturamento anual",
                    format_money(
                        get_value(
                            selected,
                            "annual_revenue",
                            default=0
                        )
                    )
                )

                col2.metric(
                    "Regime",
                    get_value(
                        selected,
                        "tax_regime",
                        default="-"
                    )
                )

                col3.metric(
                    "Setor",
                    get_value(
                        selected,
                        "sector",
                        default="-"
                    )
                )

                st.write(
                    "**Contato:**",
                    get_value(
                        selected,
                        "contact_name",
                        default="-"
                    )
                )

                st.write(
                    "**E-mail:**",
                    get_value(
                        selected,
                        "contact_email",
                        default="-"
                    )
                )

                st.write(
                    "**Telefone:**",
                    get_value(
                        selected,
                        "contact_phone",
                        default="-"
                    )
                )

                st.write(
                    "**Contador:**",
                    get_value(
                        selected,
                        "accountant",
                        default="-"
                    )
                )


            # STATUS

            with tab2:

                current_commercial = (
                    get_value(
                        selected,
                        "commercial_status",
                        default="Em prospecção"
                    )
                )

                current_operational = (
                    get_value(
                        selected,
                        "operational_status",
                        default=(
                            "Aguardando contratação"
                        )
                    )
                )

                with st.form(
                    "client_status_form"
                ):

                    commercial = (
                        st.selectbox(

                            "Status comercial",

                            COMMERCIAL_STATUSES,

                            index=(
                                COMMERCIAL_STATUSES.index(
                                    current_commercial
                                )

                                if current_commercial
                                in COMMERCIAL_STATUSES

                                else 0
                            )

                        )
                    )

                    operational = (
                        st.selectbox(

                            "Status operacional",

                            OPERATIONAL_STATUSES,

                            index=(
                                OPERATIONAL_STATUSES.index(
                                    current_operational
                                )

                                if current_operational
                                in OPERATIONAL_STATUSES

                                else 0
                            )

                        )
                    )

                    note = (
                        st.text_area(
                            "Observação da mudança"
                        )
                    )

                    update_status = (
                        st.form_submit_button(
                            "Atualizar status",
                            width="stretch"
                        )
                    )


                if update_status:

                    db_update_flexible(

                        "clients",

                        client_id,

                        {

                            "commercial_status":
                            commercial,

                            "operational_status":
                            operational,

                            "updated_at":
                            now_iso()

                        }

                    )

                    add_timeline(

                        client_id,

                        "Status atualizado",

                        (
                            f"{commercial}"
                            " | "
                            f"{operational}"
                            ". "
                            f"{note}"
                        ),

                        "Status"

                    )

                    st.success(
                        "Status atualizado."
                    )

                    st.rerun()


            # CASOS

            with tab3:

                client_cases = [

                    case
                    for case in cases()

                    if str(
                        case.get("client_id")
                    )
                    == str(client_id)

                ]

                if client_cases:

                    dataframe = pd.DataFrame([

                        {

                            "Caso":
                            get_value(
                                case,
                                "title"
                            ),

                            "Tipo":
                            get_value(
                                case,
                                "case_type"
                            ),

                            "Urgência":
                            get_value(
                                case,
                                "urgency"
                            ),

                            "Status":
                            get_value(
                                case,
                                "status"
                            )

                        }

                        for case
                        in client_cases

                    ])

                    st.dataframe(
                        dataframe,
                        hide_index=True,
                        width="stretch"
                    )

                else:

                    st.info(
                        "Nenhum caso deste cliente."
                    )


            # HISTÓRICO

            with tab4:

                timeline_table = (
                    first_existing_table(
                        "client_timeline",
                        "timeline"
                    )
                )

                if timeline_table:

                    timeline = db_select(

                        timeline_table,

                        {
                            "client_id":
                            client_id
                        },

                        order="created_at",

                        desc=True,

                        silent=True

                    )

                    if timeline:

                        for item in timeline:

                            st.markdown(

                                "**"
                                + str(
                                    get_value(
                                        item,
                                        "title",
                                        "event_type",
                                        default="Atualização"
                                    )
                                )
                                + "**"
                            )

                            st.write(
                                get_value(
                                    item,
                                    "description",
                                    "details",
                                    default=""
                                )
                            )

                            st.caption(
                                get_value(
                                    item,
                                    "created_at",
                                    default=""
                                )
                            )

                            st.divider()

                    else:

                        st.info(
                            "Sem eventos no histórico."
                        )

                else:

                    st.info(
                        "Tabela de histórico "
                        "não disponível."
                    )


# ============================================================
# CASOS
# ============================================================

elif page == "Casos":

    header(
        "Casos",
        "Checklist documental → dor → "
        "análise preliminar → diagnóstico."
    )

    case = choose_case(
        "Selecione um caso",
        key="case_page"
    )


    if case:

        case_id = case.get("id")

        client_id = case.get(
            "client_id"
        )

        client = get_client(
            client_id
        )

        st.markdown(
            "## "
            + str(
                get_value(
                    case,
                    "title",
                    default="Caso"
                )
            )
        )

        st.caption(
            f"{row_name(client or {})}"
            " | "
            f"{get_value(case, 'case_type')}"
            " | "
            f"{get_value(case, 'urgency')}"
            " | "
            f"{get_value(case, 'status')}"
        )


        (
            tab_documents,
            tab_pain,
            tab_preliminary,
            tab_diagnosis
        ) = st.tabs(

            [

                "1. Documentos",

                "2. Dor do cliente",

                "3. Análise preliminar",

                "4. Diagnóstico"

            ]

        )


        # ====================================================
        # DOCUMENTOS
        # ====================================================

        with tab_documents:

            st.markdown(
                "### Checklist documental"
            )

            existing_documents = (
                case_documents(
                    case_id
                )
            )

            uploaded_keys = {

                get_value(
                    document,
                    "rkey",
                    "category",
                    "subcategory",
                    "label"
                )

                for document
                in existing_documents

            }


            checklist_rows = []

            for (
                key,
                label,
                required,
                note
            ) in CHECKLIST:

                checklist_rows.append({

                    "Documento":
                    label,

                    "Obrigatório-base":
                    (
                        "Sim"
                        if required
                        else "Conforme o caso"
                    ),

                    "Status":
                    (
                        "Anexado"
                        if key in uploaded_keys
                        else "Pendente"
                    ),

                    "Observação":
                    note

                })


            st.dataframe(

                pd.DataFrame(
                    checklist_rows
                ),

                hide_index=True,

                width="stretch"

            )


            category_map = {

                key: (
                    label,
                    required,
                    note
                )

                for (
                    key,
                    label,
                    required,
                    note
                )
                in CHECKLIST

            }


            category = st.selectbox(

                "Tipo de documento "
                "que será anexado",

                list(
                    category_map.keys()
                ),

                format_func=lambda value:
                category_map[value][0]

            )


            uploads = st.file_uploader(

                "Selecione um ou mais arquivos",

                type=[

                    "pdf",

                    "docx",

                    "txt",

                    "csv",

                    "xlsx",

                    "xls",

                    "json",

                    "xml",

                    "png",

                    "jpg",

                    "jpeg"

                ],

                accept_multiple_files=True

            )


            save_documents = st.button(

                "Salvar documentos no cliente",

                width="stretch",

                disabled=not uploads

            )


            if save_documents:

                success = 0

                for file in uploads or []:

                    raw = (
                        file.getvalue()
                    )

                    unique_name = (

                        datetime.now()
                        .strftime(
                            "%Y%m%d_%H%M%S_%f"
                        )

                        + "_"

                        + clean_name(
                            file.name
                        )

                    )


                    # Cada empresa possui sua
                    # própria estrutura de pastas.

                    path = (

                        f"clients/{client_id}"
                        f"/cases/{case_id}"
                        f"/{category}"
                        f"/{unique_name}"

                    )


                    uploaded, error = (
                        storage_upload(

                            CLIENT_BUCKET,

                            path,

                            raw,

                            file.type

                        )
                    )


                    if not uploaded:

                        st.error(
                            f"Falha ao enviar "
                            f"{file.name}: {error}"
                        )

                        continue


                    extracted = (
                        extract_text(
                            file.name,
                            raw
                        )
                    )


                    # Salva também texto extraído.
                    if extracted:

                        storage_upload(

                            CLIENT_BUCKET,

                            sidecar_path(path),

                            extracted.encode(
                                "utf-8"
                            ),

                            (
                                "text/plain; "
                                "charset=utf-8"
                            )

                        )


                    digest = hashlib.sha256(
                        raw
                    ).hexdigest()


                    saved = db_insert_flexible(

                        "documents",

                        {

                            "client_id":
                            client_id,

                            "case_id":
                            case_id,

                            "rkey":
                            category,

                            "category":
                            category,

                            "subcategory":
                            category,

                            "label":
                            category_map[
                                category
                            ][0],

                            "original_name":
                            file.name,

                            "file_name":
                            file.name,

                            "storage_path":
                            path,

                            "stored_path":
                            path,

                            "path":
                            path,

                            "hash_sha256":
                            digest,

                            "sha256":
                            digest,

                            "status":
                            "Anexado",

                            "extracted_text":
                            (
                                extracted[
                                    :200000
                                ]
                                if extracted
                                else ""
                            ),

                            "uploaded_by":
                            st.session_state[
                                "user"
                            ][
                                "email"
                            ],

                            "created_at":
                            now_iso()

                        }

                    )


                    if saved:

                        success += 1


                if success:

                    db_update_flexible(

                        "cases",

                        case_id,

                        {

                            "status":
                            "Aguardando documentos",

                            "updated_at":
                            now_iso()

                        },

                        silent=True

                    )


                    add_timeline(

                        client_id,

                        "Documentos anexados",

                        (
                            f"{success} arquivo(s) "
                            "adicionado(s) ao caso."
                        ),

                        "Documento"

                    )


                    st.success(
                        f"{success} arquivo(s) "
                        "salvo(s) no Supabase."
                    )

                    st.rerun()


            # -----------------------------------------------
            # LISTA DE DOCUMENTOS DO CASO
            # -----------------------------------------------

            st.markdown(
                "### Arquivos deste caso"
            )

            docs_now = (
                case_documents(
                    case_id
                )
            )


            if docs_now:

                for document in docs_now:

                    column1, column2 = (
                        st.columns(
                            [5, 1]
                        )
                    )

                    with column1:

                        st.write(
                            "**"
                            + document_name(
                                document
                            )
                            + "**"
                        )

                        st.caption(
                            str(
                                get_value(
                                    document,
                                    "label",
                                    "category",
                                    "rkey"
                                )
                            )
                            + " | "
                            + str(
                                get_value(
                                    document,
                                    "created_at"
                                )
                            )
                        )

                    with column2:

                        path = (
                            document_path(
                                document
                            )
                        )

                        if path:

                            raw, error = (
                                storage_download(
                                    CLIENT_BUCKET,
                                    path
                                )
                            )

                            if raw:

                                st.download_button(

                                    "Baixar",

                                    data=raw,

                                    file_name=(
                                        document_name(
                                            document
                                        )
                                    ),

                                    key=(
                                        "download_"
                                        + str(
                                            document.get(
                                                "id"
                                            )
                                        )
                                        + "_"
                                        + hashlib.md5(
                                            path.encode()
                                        ).hexdigest()
                                    ),

                                    width="stretch"

                                )

            else:

                st.info(
                    "Nenhum documento anexado."
                )


        # ====================================================
        # DOR DO CLIENTE
        # ====================================================

        with tab_pain:

            with st.form(
                "pain_form"
            ):

                summary = st.text_input(

                    "Resumo da dor",

                    value=get_value(
                        case,
                        "pain_summary"
                    )

                )

                details = st.text_area(

                    "Relato completo",

                    value=get_value(
                        case,
                        "pain_details"
                    ),

                    height=220

                )

                objectives = st.text_area(

                    "Objetivo do cliente",

                    value=get_value(
                        case,
                        "objectives"
                    ),

                    height=100

                )


                save_pain = (
                    st.form_submit_button(
                        "Salvar relato",
                        width="stretch"
                    )
                )


            if save_pain:

                db_update_flexible(

                    "cases",

                    case_id,

                    {

                        "pain_summary":
                        summary,

                        "pain_details":
                        details,

                        "objectives":
                        objectives,

                        "updated_at":
                        now_iso()

                    }

                )


                add_timeline(

                    client_id,

                    "Dor do cliente atualizada",

                    summary,

                    "Atendimento"

                )


                st.success(
                    "Relato salvo."
                )

                st.rerun()


        # ====================================================
        # ANÁLISE PRELIMINAR
        # ====================================================

        with tab_preliminary:

            st.info(

                "A análise preliminar verifica "
                "o material anexado, separa fatos "
                "de alegações, identifica documentos "
                "faltantes e aponta riscos urgentes."

            )


            run_preliminary = (
                st.button(

                    "Executar análise preliminar "
                    "com OpenAI",

                    type="primary",

                    width="stretch"

                )
            )


            if run_preliminary:

                with st.spinner(
                    "Lendo documentos "
                    "e realizando análise..."
                ):

                    current_case = (
                        get_case(case_id)
                        or case
                    )

                    context = (
                        build_case_context(
                            case_id
                        )
                    )


                    prompt = f"""

CLIENTE:
{row_name(client or {})}

CNPJ:
{get_value(client or {}, "cnpj")}

REGIME:
{get_value(client or {}, "tax_regime")}

CASO:
{get_value(current_case, "title")}

TIPO:
{get_value(current_case, "case_type")}

URGÊNCIA:
{get_value(current_case, "urgency")}


DOR RELATADA:

{get_value(
    current_case,
    "pain_details",
    "pain_summary"
)}


OBJETIVO:

{get_value(
    current_case,
    "objectives"
)}


DOCUMENTOS EXTRAÍDOS:

{
    context
    if context
    else "[Nenhum texto documental extraído.]"
}


TAREFA:

Produza uma ANÁLISE PRELIMINAR,
sem fechar diagnóstico definitivo.

Estruture obrigatoriamente em:

1. RESUMO OBJETIVO DO CASO.

2. FATOS EFETIVAMENTE COMPROVADOS
PELOS DOCUMENTOS.

3. INFORMAÇÕES APENAS RELATADAS
E AINDA NÃO COMPROVADAS.

4. DOCUMENTOS RECEBIDOS E
UTILIDADE DE CADA UM.

5. DOCUMENTOS OU INFORMAÇÕES
ADICIONAIS NECESSÁRIOS,
EXPLICANDO O MOTIVO.

6. PRAZOS, DATAS DE CIÊNCIA
E RISCOS URGENTES.

7. POSSÍVEIS FRENTES
ADMINISTRATIVAS.

8. POSSÍVEIS FRENTES JUDICIAIS,
APENAS PRELIMINARMENTE.

9. POSSÍVEIS CRÉDITOS TRIBUTÁRIOS
A INVESTIGAR.

Não trate valor teórico como
crédito validado.

10. PRÓXIMOS PASSOS EM ORDEM
DE PRIORIDADE.

"""


                    try:

                        text, model = (
                            openai_response(
                                prompt,
                                web_search=False
                            )
                        )


                        save_case_ai(

                            case_id,

                            client_id,

                            "preliminary",

                            text,

                            model

                        )


                        db_update_flexible(

                            "clients",

                            client_id,

                            {

                                "commercial_status":
                                "Diagnóstico em andamento",

                                "updated_at":
                                now_iso()

                            },

                            silent=True

                        )


                        st.success(
                            "Análise concluída "
                            f"com {model}."
                        )

                        st.rerun()


                    except Exception as error:

                        st.error(
                            "Falha na OpenAI: "
                            f"{error}"
                        )


            current_case = (
                get_case(case_id)
                or case
            )

            preliminary = (
                get_value(
                    current_case,
                    "ai_preliminary"
                )
            )


            if preliminary:

                st.markdown(
                    preliminary
                )

                document = (
                    text_to_docx(
                        "Análise Preliminar Tributária",
                        preliminary
                    )
                )

                st.download_button(

                    "Baixar análise em DOCX",

                    document,

                    file_name=(
                        "analise_preliminar_"
                        + str(case_id)[:8]
                        + ".docx"
                    ),

                    mime=(
                        "application/vnd.openxmlformats-"
                        "officedocument.wordprocessingml."
                        "document"
                    ),

                    width="stretch"

                )


        # ====================================================
        # DIAGNÓSTICO
        # ====================================================

        with tab_diagnosis:

            st.warning(

                "O diagnóstico utiliza os documentos "
                "do caso, a base jurídica interna "
                "e, quando habilitada, pesquisa web. "
                "A revisão profissional continua "
                "obrigatória."

            )


            use_web = st.checkbox(

                "Pesquisar legislação e "
                "jurisprudência atual na web",

                value=True

            )


            run_diagnosis = (
                st.button(

                    "Gerar diagnóstico "
                    "tributário completo",

                    type="primary",

                    width="stretch"

                )
            )


            if run_diagnosis:

                with st.spinner(
                    "Gerando diagnóstico técnico..."
                ):

                    current_case = (
                        get_case(case_id)
                        or case
                    )

                    context = (
                        build_case_context(
                            case_id
                        )
                    )

                    legal_context = (
                        load_legal_context()
                    )

                    preliminary = (
                        get_value(
                            current_case,
                            "ai_preliminary"
                        )
                    )


                    prompt = f"""

CLIENTE:
{row_name(client or {})}

CNPJ:
{get_value(client or {}, "cnpj")}

REGIME:
{get_value(client or {}, "tax_regime")}

CASO:
{get_value(current_case, "title")}

TIPO:
{get_value(current_case, "case_type")}

URGÊNCIA:
{get_value(current_case, "urgency")}


DOR DO CLIENTE:

{get_value(
    current_case,
    "pain_details",
    "pain_summary"
)}


OBJETIVO:

{get_value(
    current_case,
    "objectives"
)}


ANÁLISE PRELIMINAR:

{
    preliminary
    or "[Ainda não produzida]"
}


DOCUMENTOS DO CASO:

{
    context
    if context
    else "[Sem texto documental extraído]"
}


BASE JURÍDICA INTERNA:

{
    legal_context
    if legal_context
    else "[Base jurídica interna vazia]"
}


TAREFA:

Produza um DIAGNÓSTICO TRIBUTÁRIO
completo, conservador e auditável.

Estruture obrigatoriamente em:

A. SUMÁRIO EXECUTIVO.

B. LINHA DO TEMPO DOS FATOS
E DATAS RELEVANTES.

C. MATRIZ DE DOCUMENTOS E PROVAS:
- documento;
- fato que comprova;
- relevância;
- lacunas.

D. ENQUADRAMENTO TRIBUTÁRIO
E PROCESSUAL.

E. REGULARIDADE FISCAL E CERTIDÕES,
SE PERTINENTE.

F. RECUPERAÇÃO DE CRÉDITOS:

Para cada possível crédito:
- tributo;
- período;
- origem;
- fundamento;
- documentação;
- condição para aproveitamento;
- risco;
- procedimento;
- necessidade de validação contábil.

Diferencie claramente:
- crédito potencial;
- crédito identificado;
- crédito validado;
- crédito aproveitável;
- benefício realizado.

G. ESTRATÉGIA ADMINISTRATIVA:

Indique:
- medida;
- finalidade;
- órgão;
- canal provável;
- documentos;
- fundamento;
- prazo;
- sequência;
- risco;
- alternativa.

H. ESTRATÉGIA JUDICIAL:

Somente proponha judicialização
se houver fato concreto e
adequação processual.

Analise:
- medida possível;
- cabimento;
- competência;
- legitimidade;
- autoridade ou polo passivo;
- prova;
- urgência;
- risco;
- necessidade de perícia;
- alternativa administrativa.

I. CONTRATOS PÚBLICOS E
REEQUILÍBRIO, SE HOUVER.

Avalie:
- fato gerador do desequilíbrio;
- matriz de riscos;
- impacto líquido;
- memória de cálculo;
- documentação;
- procedimento administrativo;
- eventual judicialização.

J. PLANO DE AÇÃO:

Divida em:
- próximas 24 horas;
- 7 dias;
- 30 dias;
- 90 dias.

K. DOCUMENTOS AINDA FALTANTES.

L. PONTOS QUE EXIGEM
VALIDAÇÃO DO ADVOGADO.

M. PONTOS QUE EXIGEM
VALIDAÇÃO DO CONTADOR.

N. FONTES NORMATIVAS,
ADMINISTRATIVAS E
JURISPRUDENCIAIS EFETIVAMENTE
UTILIZADAS.

Não invente jurisprudência.

Se não localizar fonte segura,
declare expressamente que
a verificação permanece pendente.

"""


                    try:

                        text, model = (
                            openai_response(
                                prompt,
                                web_search=use_web
                            )
                        )


                        save_case_ai(

                            case_id,

                            client_id,

                            "diagnosis",

                            text,

                            model

                        )


                        st.success(
                            "Diagnóstico concluído "
                            f"com {model}."
                        )

                        st.rerun()


                    except Exception as error:

                        st.error(
                            "Falha na OpenAI: "
                            f"{error}"
                        )


            current_case = (
                get_case(case_id)
                or case
            )

            diagnosis = (
                get_value(
                    current_case,
                    "ai_diagnosis"
                )
            )


            if diagnosis:

                st.markdown(
                    diagnosis
                )

                document = (
                    text_to_docx(
                        "Diagnóstico Tributário",
                        diagnosis
                    )
                )

                st.download_button(

                    "Baixar diagnóstico em DOCX",

                    document,

                    file_name=(
                        "diagnostico_tributario_"
                        + str(case_id)[:8]
                        + ".docx"
                    ),

                    mime=(
                        "application/vnd.openxmlformats-"
                        "officedocument.wordprocessingml."
                        "document"
                    ),

                    width="stretch"

                )


# ============================================================
# TAREFAS E PRAZOS
# ============================================================

elif page == "Tarefas e prazos":

    header(
        "Tarefas e prazos",
        "Controle operacional e jurídico."
    )


    if not db_table_exists(
        "tasks"
    ):

        st.warning(
            "A tabela tasks não está "
            "disponível no Supabase."
        )

    else:

        client = choose_client(
            "Cliente",
            key="task_client"
        )


        if client:

            with st.form(
                "new_task"
            ):

                title = st.text_input(
                    "Tarefa"
                )

                description = st.text_area(
                    "Descrição"
                )

                col1, col2, col3 = (
                    st.columns(3)
                )

                responsible = (
                    col1.text_input(
                        "Responsável",
                        value="Rafael"
                    )
                )

                priority = (
                    col2.selectbox(
                        "Prioridade",
                        [
                            "Baixa",
                            "Média",
                            "Alta",
                            "Crítica"
                        ]
                    )
                )

                due = (
                    col3.date_input(
                        "Prazo",
                        value=date.today()
                    )
                )

                status = st.selectbox(

                    "Status",

                    [

                        "Aberta",

                        "Em andamento",

                        "Aguardando cliente",

                        "Aguardando contador",

                        "Concluída",

                        "Cancelada"

                    ]

                )


                create_task = (
                    st.form_submit_button(
                        "Criar tarefa",
                        width="stretch"
                    )
                )


            if create_task:

                db_insert_flexible(

                    "tasks",

                    {

                        "client_id":
                        client.get("id"),

                        "title":
                        title,

                        "description":
                        description,

                        "responsible":
                        responsible,

                        "priority":
                        priority,

                        "due_date":
                        due.isoformat(),

                        "status":
                        status,

                        "created_at":
                        now_iso()

                    }

                )


                add_timeline(

                    client.get("id"),

                    "Tarefa criada",

                    title,

                    "Tarefa"

                )


                st.success(
                    "Tarefa criada."
                )

                st.rerun()


        task_rows = db_select(

            "tasks",

            order="due_date",

            desc=False,

            silent=True

        )


        if task_rows:

            st.dataframe(
                pd.DataFrame(task_rows),
                hide_index=True,
                width="stretch"
            )


# ============================================================
# DOCUMENTOS
# ============================================================

elif page == "Documentos":

    header(
        "Documentos",
        "Data room central dos clientes."
    )

    document_rows = db_select(

        "documents",

        order="created_at",

        desc=True,

        silent=True

    )

    clients_by_id = (
        client_map()
    )


    if document_rows:

        rows = []

        for document in document_rows:

            client = clients_by_id.get(

                str(
                    document.get(
                        "client_id"
                    )
                ),

                {}

            )


            rows.append({

                "Cliente":
                row_name(client),

                "Arquivo":
                document_name(
                    document
                ),

                "Categoria":
                get_value(
                    document,
                    "label",
                    "category",
                    "rkey"
                ),

                "Caso":
                str(
                    document.get(
                        "case_id"
                    )
                    or ""
                ),

                "Data":
                get_value(
                    document,
                    "created_at"
                )

            })


        st.dataframe(
            pd.DataFrame(rows),
            hide_index=True,
            width="stretch"
        )


    else:

        st.info(
            "Nenhum documento cadastrado."
        )


# ============================================================
# BASE JURÍDICA
# ============================================================

elif page == "Base jurídica":

    header(

        "Base jurídica",

        "Legislação, regulamentos, notas técnicas, "
        "decisões e materiais de apoio da IA."

    )


    legal_uploads = st.file_uploader(

        "Adicionar arquivos à base jurídica",

        type=[

            "pdf",

            "docx",

            "txt",

            "md",

            "csv",

            "xlsx",

            "xls",

            "json",

            "xml"

        ],

        accept_multiple_files=True,

        key="legal_upload"

    )


    save_legal = st.button(

        "Salvar na base jurídica",

        width="stretch",

        disabled=not legal_uploads

    )


    if save_legal:

        total = 0

        for file in legal_uploads or []:

            raw = file.getvalue()

            unique_name = (

                datetime.now()
                .strftime(
                    "%Y%m%d_%H%M%S_%f"
                )

                + "_"

                + clean_name(
                    file.name
                )

            )

            path = (
                f"library/{unique_name}"
            )


            uploaded, error = (
                storage_upload(

                    LEGAL_BUCKET,

                    path,

                    raw,

                    file.type

                )
            )


            if not uploaded:

                st.error(
                    f"{file.name}: {error}"
                )

                continue


            extracted = extract_text(
                file.name,
                raw
            )


            if extracted:

                side_name = (

                    unique_name
                    .rsplit(".", 1)[0]
                    + ".txt"

                )


                storage_upload(

                    LEGAL_BUCKET,

                    f"_text/{side_name}",

                    (
                        "ARQUIVO ORIGINAL: "
                        f"{file.name}"
                        "\n\n"
                        f"{extracted}"
                    ).encode(
                        "utf-8"
                    ),

                    (
                        "text/plain; "
                        "charset=utf-8"
                    )

                )


            total += 1


        if total:

            st.success(
                f"{total} arquivo(s) "
                "adicionado(s) à base."
            )


    # LISTA

    st.subheader(
        "Arquivos cadastrados"
    )

    legal_files = storage_list(
        LEGAL_BUCKET,
        "library"
    )


    if legal_files:

        for item in legal_files:

            name = item.get(
                "name"
            )

            if name:

                st.write(
                    "•",
                    name
                )

    else:

        st.info(
            "A base jurídica está vazia."
        )


    # PESQUISA

    st.subheader(
        "Pesquisar com IA"
    )

    question = st.text_area(

        "Pergunta",

        placeholder=(
            "Ex.: Quais normas da minha "
            "base tratam de compensação "
            "tributária e PER/DCOMP?"
        ),

        height=100

    )


    web = st.checkbox(

        "Complementar com "
        "pesquisa web oficial",

        value=True,

        key="legal_web"

    )


    search_legal = st.button(

        "Pesquisar base jurídica",

        type="primary",

        width="stretch",

        disabled=not question.strip()

    )


    if search_legal:

        with st.spinner(
            "Pesquisando..."
        ):

            base = load_legal_context(
                max_chars=100000
            )


            prompt = f"""

PERGUNTA:

{question}


BASE JURÍDICA INTERNA:

{
    base
    if base
    else "[Base interna vazia]"
}


Responda utilizando primeiro
a base interna.

Identifique os arquivos internos
efetivamente utilizados.

Se a pesquisa web estiver habilitada,
complemente com fontes oficiais.

Separe a resposta em:

1. Resposta objetiva.

2. Documentos internos utilizados.

3. Legislação.

4. Jurisprudência, se houver.

5. Fontes externas.

6. Pontos ainda incertos.

"""


            try:

                answer, model = (
                    openai_response(
                        prompt,
                        web_search=web
                    )
                )

                st.caption(
                    f"Modelo utilizado: {model}"
                )

                st.markdown(
                    answer
                )

            except Exception as error:

                st.error(
                    "Falha na OpenAI: "
                    f"{error}"
                )


# ============================================================
# PROPOSTAS E CONTRATOS
# ============================================================

elif page == "Propostas e contratos":

    header(

        "Propostas e contratos",

        "Do diagnóstico à contratação "
        "e execução."

    )


    client = choose_client(
        "Cliente",
        key="proposal_client"
    )


    tab_proposal, tab_contract = (
        st.tabs(
            [
                "Propostas",
                "Contratos"
            ]
        )
    )


    # --------------------------------------------------------
    # PROPOSTAS
    # --------------------------------------------------------

    with tab_proposal:

        if client:

            with st.form(
                "proposal_form"
            ):

                title = st.text_input(

                    "Objeto da proposta",

                    value=(
                        "Diagnóstico Tributário "
                        "e Plano de Ação"
                    )

                )

                amount = st.number_input(

                    "Honorários fixos",

                    min_value=0.0,

                    step=1000.0

                )

                success_fee = (
                    st.number_input(

                        "Êxito (%)",

                        min_value=0.0,

                        max_value=100.0,

                        step=1.0

                    )
                )

                scope = st.text_area(

                    "Escopo",

                    height=150

                )

                status = st.selectbox(

                    "Status",

                    [

                        "Rascunho",

                        "Enviada",

                        "Em negociação",

                        "Aceita",

                        "Recusada"

                    ]

                )

                create_proposal = (
                    st.form_submit_button(

                        "Salvar proposta",

                        width="stretch"

                    )
                )


            if create_proposal:

                table = (
                    first_existing_table(

                        "proposals",

                        "commercial_proposals"

                    )
                )


                if table:

                    db_insert_flexible(

                        table,

                        {

                            "client_id":
                            client.get("id"),

                            "title":
                            title,

                            "amount":
                            amount,

                            "fixed_fee":
                            amount,

                            "success_fee_percent":
                            success_fee,

                            "scope":
                            scope,

                            "status":
                            status,

                            "created_at":
                            now_iso(),

                            "updated_at":
                            now_iso()

                        }

                    )


                    db_update_flexible(

                        "clients",

                        client.get("id"),

                        {

                            "commercial_status":
                            "Proposta enviada",

                            "updated_at":
                            now_iso()

                        },

                        silent=True

                    )


                    add_timeline(

                        client.get("id"),

                        "Proposta registrada",

                        (
                            title
                            + " — "
                            + format_money(
                                amount
                            )
                        ),

                        "Comercial"

                    )


                    st.success(
                        "Proposta salva."
                    )


                else:

                    st.error(
                        "Tabela de propostas "
                        "não encontrada."
                    )


        proposal_table = (
            first_existing_table(

                "proposals",

                "commercial_proposals"

            )
        )


        if proposal_table:

            rows = db_select(

                proposal_table,

                order="created_at",

                desc=True,

                silent=True

            )

            if rows:

                st.dataframe(

                    pd.DataFrame(rows),

                    hide_index=True,

                    width="stretch"

                )


    # --------------------------------------------------------
    # CONTRATOS
    # --------------------------------------------------------

    with tab_contract:

        if client:

            st.info(
                "Registre aqui a contratação "
                "e acompanhe a execução."
            )


            with st.form(
                "contract_form"
            ):

                title = st.text_input(
                    "Objeto do contrato"
                )

                value = st.number_input(

                    "Valor do contrato/honorários",

                    min_value=0.0,

                    step=1000.0

                )

                status = st.selectbox(

                    "Status",

                    [

                        "Minuta",

                        "Enviado para assinatura",

                        "Assinado",

                        "Em execução",

                        "Executado",

                        "Finalizado",

                        "Rescindido"

                    ]

                )

                notes = st.text_area(
                    "Observações"
                )

                create_contract = (
                    st.form_submit_button(

                        "Salvar contrato",

                        width="stretch"

                    )
                )


            if create_contract:

                table = first_existing_table(

                    "contracts",

                    "legal_contracts"

                )


                if table:

                    db_insert_flexible(

                        table,

                        {

                            "client_id":
                            client.get("id"),

                            "title":
                            title,

                            "object":
                            title,

                            "amount":
                            value,

                            "value":
                            value,

                            "status":
                            status,

                            "notes":
                            notes,

                            "created_at":
                            now_iso(),

                            "updated_at":
                            now_iso()

                        }

                    )


                    if status in (
                        "Minuta",
                        "Enviado para assinatura"
                    ):

                        commercial = (
                            "Proposta enviada"
                        )

                    else:

                        commercial = (
                            "Contratado"
                        )


                    if status == "Em execução":

                        operational = (
                            "Em execução"
                        )

                    elif status in (
                        "Executado",
                        "Finalizado"
                    ):

                        operational = status

                    else:

                        operational = (
                            "Aguardando início "
                            "da execução"
                        )


                    db_update_flexible(

                        "clients",

                        client.get("id"),

                        {

                            "commercial_status":
                            commercial,

                            "operational_status":
                            operational,

                            "updated_at":
                            now_iso()

                        },

                        silent=True

                    )


                    add_timeline(

                        client.get("id"),

                        "Contrato atualizado",

                        (
                            title
                            + " — "
                            + status
                        ),

                        "Contrato"

                    )


                    st.success(
                        "Contrato salvo."
                    )


                else:

                    st.error(
                        "Tabela de contratos "
                        "não encontrada."
                    )


        contract_table = (
            first_existing_table(

                "contracts",

                "legal_contracts"

            )
        )


        if contract_table:

            rows = db_select(

                contract_table,

                order="created_at",

                desc=True,

                silent=True

            )

            if rows:

                st.dataframe(

                    pd.DataFrame(rows),

                    hide_index=True,

                    width="stretch"

                )


# ============================================================
# RELATÓRIOS
# ============================================================

elif page == "Relatórios":

    header(

        "Relatórios",

        "Dossiê executivo consolidado "
        "por cliente."

    )


    client = choose_client(
        "Cliente",
        key="report_client"
    )


    if client:

        client_id = (
            client.get("id")
        )


        client_cases = [

            case
            for case in cases()

            if str(
                case.get("client_id")
            )
            == str(client_id)

        ]


        documents = [

            document

            for document
            in db_select(
                "documents",
                silent=True
            )

            if str(
                document.get(
                    "client_id"
                )
            )
            == str(client_id)

        ]


        tasks = [

            task

            for task
            in db_select(
                "tasks",
                silent=True
            )

            if str(
                task.get(
                    "client_id"
                )
            )
            == str(client_id)

        ]


        report = [

            (
                "# Dossiê Executivo — "
                + row_name(client)
            ),

            "",

            (
                "CNPJ: "
                + str(
                    get_value(
                        client,
                        "cnpj",
                        default="-"
                    )
                )
            ),

            (
                "Regime: "
                + str(
                    get_value(
                        client,
                        "tax_regime",
                        default="-"
                    )
                )
            ),

            (
                "Status comercial: "
                + str(
                    get_value(
                        client,
                        "commercial_status",
                        "status",
                        default="-"
                    )
                )
            ),

            (
                "Status operacional: "
                + str(
                    get_value(
                        client,
                        "operational_status",
                        default="-"
                    )
                )
            ),

            "",

            "## Casos"

        ]


        for case in client_cases:

            report.append(

                "- "
                + str(
                    get_value(
                        case,
                        "title"
                    )
                )

                + " | "

                + str(
                    get_value(
                        case,
                        "status"
                    )
                )

                + " | "

                + str(
                    get_value(
                        case,
                        "urgency"
                    )
                )

            )


            diagnosis = get_value(
                case,
                "ai_diagnosis"
            )


            if diagnosis:

                report.extend(

                    [

                        "",

                        (
                            "### Diagnóstico — "
                            + str(
                                get_value(
                                    case,
                                    "title"
                                )
                            )
                        ),

                        diagnosis[:20000]

                    ]

                )


        report.extend(
            [
                "",
                "## Documentos"
            ]
        )


        for document in documents:

            report.append(

                "- "
                + document_name(
                    document
                )

                + " | "

                + str(
                    get_value(
                        document,
                        "label",
                        "category",
                        "rkey"
                    )
                )

            )


        report.extend(
            [
                "",
                "## Tarefas"
            ]
        )


        for task in tasks:

            report.append(

                "- "
                + str(
                    get_value(
                        task,
                        "title"
                    )
                )

                + " | "

                + str(
                    get_value(
                        task,
                        "due_date"
                    )
                )

                + " | "

                + str(
                    get_value(
                        task,
                        "status"
                    )
                )

            )


        report_text = "\n".join(
            report
        )


        st.markdown(
            report_text
        )


        document = text_to_docx(

            (
                "Dossiê Executivo — "
                + row_name(client)
            ),

            report_text

        )


        st.download_button(

            "Baixar dossiê em DOCX",

            document,

            file_name=(

                "dossie_"

                + clean_name(
                    row_name(client)
                )

                + ".docx"

            ),

            mime=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            ),

            width="stretch"

        )


# ============================================================
# CONFIGURAÇÕES
# ============================================================

elif page == "Configurações":

    header(

        "Configurações",

        "Conectividade e diagnóstico "
        "técnico da plataforma."

    )


    col1, col2, col3 = (
        st.columns(3)
    )


    col1.metric(

        "OpenAI",

        (
            "Configurada"
            if OPENAI_API_KEY
            else "Pendente"
        )

    )


    col2.metric(

        "Supabase",

        (
            "Configurado"

            if (
                SUPABASE_URL
                and SUPABASE_SECRET_KEY
            )

            else "Pendente"
        )

    )


    col3.metric(
        "Modelo",
        OPENAI_MODEL
    )


    # TESTE SUPABASE

    if st.button(
        "Testar Supabase",
        width="stretch"
    ):

        try:

            rows = db_select(
                "clients",
                silent=True
            )

            st.success(
                "Supabase conectado. "
                "Tabela clients acessível. "
                f"Registros atuais: {len(rows)}."
            )

        except Exception as error:

            st.error(
                str(error)
            )


    # TESTE OPENAI

    if st.button(
        "Testar OpenAI",
        width="stretch"
    ):

        try:

            text, model = (
                openai_response(

                    (
                        "Responda somente: "
                        "CONEXÃO OPENAI OK"
                    ),

                    web_search=False

                )
            )

            st.success(
                f"{model}: {text}"
            )

        except Exception as error:

            st.error(
                str(error)
            )


    # BUCKETS

    st.subheader(
        "Buckets"
    )

    st.code(

        f'SUPABASE_CLIENT_BUCKET = "{CLIENT_BUCKET}"\n'

        f'SUPABASE_LEGAL_BUCKET = "{LEGAL_BUCKET}"'

    )


    # TABELAS

    st.subheader(
        "Tabelas detectadas"
    )


    candidate_tables = [

        "clients",

        "cases",

        "documents",

        "tasks",

        "client_timeline",

        "proposals",

        "contracts",

        "diagnostics",

        "generated_documents"

    ]


    detected = [

        {

            "Tabela":
            table,

            "Disponível":
            (
                "Sim"
                if db_table_exists(table)
                else "Não"
            )

        }

        for table
        in candidate_tables

    ]


    st.dataframe(

        pd.DataFrame(
            detected
        ),

        hide_index=True,

        width="stretch"

    )


    st.warning(

        "MVP de validação. Antes do uso profissional "
        "com dados fiscais sensíveis, altere a senha "
        "padrão, configure autenticação robusta, "
        "revise RLS, políticas de acesso, backups, "
        "auditoria, retenção documental e LGPD."

    )


    st.info(

        "A IA analisa documentos, pesquisa fontes, "
        "gera diagnósticos e minutas. "
        "O sistema não protocola automaticamente "
        "no e-CAC, PJe ou SEI sem integração "
        "autorizada, credenciais e certificado digital."

    )